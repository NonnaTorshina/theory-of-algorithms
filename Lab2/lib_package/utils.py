from docx import Document
from openpyxl import Workbook
import json
from typing import Dict, Any


# Генератор отчётов
class ReportGenerator:
    @staticmethod
    # Сохраняет в формате docx
    def save_to_docx(data: dict, filename: str):

        doc = Document()
        doc.add_heading('Отчет электронной библиотеки', 0)

        # Статистика читателя
        doc.add_heading('Статистика читателя', level=1)
        reader_stats = data['reader_statistics']
        doc.add_paragraph(f"Читатель: {reader_stats['reader_name']}")
        doc.add_paragraph(f"Прочитано книг: {reader_stats['total_books_read']}")
        doc.add_paragraph(f"Книг в процессе: {reader_stats['books_in_progress']}")
        doc.add_paragraph(f"Всего страниц: {reader_stats['total_pages_read']}")
        doc.add_paragraph(f"Рейтинг чтения: {reader_stats['average_rating']}")

        # Статистика абонемента
        doc.add_heading('Статистика абонемента', level=1)
        sub_stats = data['subscription_statistics']
        doc.add_paragraph(f"Статус: {sub_stats['subscription_status']}")
        doc.add_paragraph(f"Дней осталось: {sub_stats['days_remaining']}")
        doc.add_paragraph(f"Дата окончания: {sub_stats['end_date']}")

        doc.save(filename)
        print(f"Отчет сохранен в {filename}")

    @staticmethod
    #Отчет в формате xlsx
    def save_to_xlsx(data: dict, filename: str):
        try:
            # Добавляем timestamp к имени файла чтобы избежать конфликтов
            import time
            timestamp = int(time.time())
            filename = f"Library_report.xlsx"

            wb = Workbook()
            ws = wb.active
            ws.title = "Отчет библиотеки"

            # Заголовки
            ws['A1'] = 'Электронная библиотека - Отчет'
            ws['A3'] = 'Статистика читателя'

            # Данные читателя
            reader_stats = data['reader_statistics']
            ws['A4'] = 'Читатель'
            ws['B4'] = reader_stats['reader_name']
            ws['A5'] = 'Прочитано книг'
            ws['B5'] = reader_stats['total_books_read']
            ws['A6'] = 'Книг в процессе'
            ws['B6'] = reader_stats['books_in_progress']
            ws['A7'] = 'Всего страниц'
            ws['B7'] = reader_stats['total_pages_read']
            ws['A8'] = 'Рейтинг чтения'
            ws['B8'] = reader_stats['average_rating']

            # Данные абонемента
            ws['A10'] = 'Статистика абонемента'
            sub_stats = data['subscription_statistics']
            ws['A11'] = 'Статус'
            ws['B11'] = sub_stats['subscription_status']
            ws['A12'] = 'Дней осталось'
            ws['B12'] = sub_stats['days_remaining']
            ws['A13'] = 'Дата окончания'
            ws['B13'] = sub_stats['end_date']

            wb.save(filename)
            print(f"Отчет сохранен в {filename}")

        except Exception as e:
            print(f"Ошибка при сохранении XLSX: {e}")

    @staticmethod
    def save_to_json(data: Dict[str, Any], filename: str):
        """Сохраняет отчет в формате JSON"""
        try:
            with open(filename, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            print(f"Отчет сохранен в {filename}")
        except Exception as e:
            print(f"Ошибка при сохранении JSON: {e}")

    @staticmethod
    #Отчет в формате txt
    def save_to_txt(data: Dict[str, Any], filename: str):
        try:
            with open(filename, 'w', encoding='utf-8') as f:
                f.write("ОТЧЕТ ЭЛЕКТРОННОЙ БИБЛИОТЕКИ\n")
                f.write("=" * 50 + "\n\n")

                f.write(f"Дата формирования: {data.get('report_date', 'Неизвестно')}\n\n")

                # Статистика читателя
                f.write("СТАТИСТИКА ЧИТАТЕЛЯ:\n")
                f.write("-" * 30 + "\n")
                reader_stats = data['reader_statistics']
                f.write(f"Читатель: {reader_stats['reader_name']}\n")
                f.write(f"Прочитано книг: {reader_stats['total_books_read']}\n")
                f.write(f"Книг в процессе: {reader_stats['books_in_progress']}\n")
                f.write(f"Всего страниц: {reader_stats['total_pages_read']}\n")
                f.write(f"Средний рейтинг: {reader_stats['average_rating']}%\n\n")

                # Прогресс по книгам
                if reader_stats['book_ratings']:
                    f.write("ПРОГРЕСС ПО КНИГАМ:\n")
                    f.write("-" * 30 + "\n")
                    for rating in reader_stats['book_ratings']:
                        f.write(f"{rating['book_name']}: {rating['progress_percent']}% ")
                        f.write(f"({rating['pages_read']}/{rating['total_pages']} стр.)\n")
                    f.write("\n")

                # Статистика абонемента
                f.write("СТАТИСТИКА АБОНЕМЕНТА:\n")
                f.write("-" * 30 + "\n")
                sub_stats = data['subscription_statistics']
                f.write(f"Статус: {sub_stats['subscription_status']}\n")
                f.write(f"Дней осталось: {sub_stats['days_remaining']}\n")
                f.write(f"Дата окончания: {sub_stats['end_date']}\n")

            print(f"Отчет сохранен в {filename}")
        except Exception as e:
            print(f"Ошибка при сохранении TXT: {e}")